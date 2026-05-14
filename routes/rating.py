from flask import Blueprint, render_template, redirect, url_for, flash, request, send_file, jsonify
from flask_jwt_extended import jwt_required, get_jwt_identity
from extensions import db
from decorators.roles import roles_required
from models.user import User
from models.award import Award, Book, JournalArticle, CollectionArticle, Dissertation, Abstract, Internet, Conference, Training, RatingTemplate, EntityCoauthor, EntitySupervisor
from forms.rating import AwardForm, PublicationForm, ConferenceForm, TrainingForm, SearchFilterForm, PUBLICATION_TYPES, AWARD_TYPES, CONFERENCE_ROLES, AWARD_LEVELS, DEGREES, FIELDS_OF_STUDY
from sqlalchemy import or_
from wtforms.validators import Optional
from datetime import datetime, timedelta, date
from io import BytesIO
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from utils import generate_gost_string


def apply_date_filter(query, date_col, date_range, date_from, date_to):
    if date_range == '3':
        cutoff = datetime.now() - timedelta(days=90)
        query = query.filter(date_col >= cutoff.date())
    elif date_range == '6':
        cutoff = datetime.now() - timedelta(days=180)
        query = query.filter(date_col >= cutoff.date())
    elif date_range == '12':
        cutoff = datetime.now() - timedelta(days=365)
        query = query.filter(date_col >= cutoff.date())
    elif date_range == '36':
        cutoff = datetime.now() - timedelta(days=1095)
        query = query.filter(date_col >= cutoff.date())
    elif date_range == '60':
        cutoff = datetime.now() - timedelta(days=1825)
        query = query.filter(date_col >= cutoff.date())
    elif date_range == 'academic_year':
        now = datetime.now()
        if now.month >= 9:
            start = date(now.year, 9, 1)
            end = date(now.year + 1, 8, 31)
        else:
            start = date(now.year - 1, 9, 1)
            end = date(now.year, 8, 31)
        query = query.filter(date_col >= start, date_col <= end)
    elif date_range == 'custom':
        if date_from:
            query = query.filter(date_col >= date_from)
        if date_to:
            query = query.filter(date_col <= date_to)
    return query


def parse_date_args(date_from_str, date_to_str):
    date_from = None
    date_to = None
    if date_from_str:
        try:
            date_from = datetime.strptime(date_from_str, '%Y-%m-%d').date()
        except ValueError:
            pass
    if date_to_str:
        try:
            date_to = datetime.strptime(date_to_str, '%Y-%m-%d').date()
        except ValueError:
            pass
    return date_from, date_to


PUB_MODELS = {
    'book': Book,
    'journal_article': JournalArticle,
    'collection_article': CollectionArticle,
    'dissertation': Dissertation,
    'abstract': Abstract,
    'internet': Internet,
}

PUB_TYPE_LABELS = {
    'book': 'Книга',
    'journal_article': 'Журнал',
    'collection_article': 'Сборник научных статей',
    'dissertation': 'Диссертация',
    'abstract': 'Автореферат',
    'internet': 'Интернет',
}


rating_bp = Blueprint('rating', __name__, url_prefix='/rating')


def filter_by_role(model, user, coauthor_entity_type=None):
    if user.role.name == 'Документовед':
        return True
    if coauthor_entity_type:
        coauthor_ids = [
            c.entity_id for c in EntityCoauthor.query.filter_by(
                entity_type=coauthor_entity_type, user_id=user.id
            ).all()
        ]
        return or_(model.user_id == user.id, model.id.in_(coauthor_ids))
    return model.user_id == user.id


# ==================== AWARDS ====================

@rating_bp.route('/awards')
@jwt_required()
def awards_list():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    search_form = SearchFilterForm()
    page = request.args.get('page', 1, type=int)

    users_list = User.query.filter(User.is_active == True).order_by(User.name).all()

    query = Award.query.filter(filter_by_role(Award, user)).filter(Award.status == 'active')

    author_ids = request.args.getlist('author_ids', type=int)
    if author_ids:
        query = query.filter(Award.user_id.in_(author_ids))

    search_query = request.args.get('search_query', '')
    if search_query:
        query = query.filter(or_(
            Award.title.ilike(f'%{search_query}%'),
            Award.description.ilike(f'%{search_query}%')
        ))

    date_range = request.args.get('date_range', 'all')
    date_from, date_to = parse_date_args(request.args.get('date_from'), request.args.get('date_to'))
    query = apply_date_filter(query, Award.date_received, date_range, date_from, date_to)

    sort_column = request.args.get('sort_column', 'date_received')
    sort_dir = request.args.get('sort_dir', 'desc')
    col_map = {
        'title': Award.title,
        'award_type': Award.award_type,
        'issuer': Award.issuer,
        'date_received': Award.date_received,
        'level': Award.level,
    }
    col = col_map.get(sort_column, Award.date_received)
    query = query.order_by(col.desc() if sort_dir == 'desc' else col.asc())

    awards = query.paginate(page=page, per_page=10)

    return render_template('rating/awards.html', awards=awards, search_form=search_form, users_list=users_list, author_ids=author_ids)

@rating_bp.route('/awards/add', methods=['GET', 'POST'])
@jwt_required()
def add_award():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    form = AwardForm()
    all_active_users = User.query.filter(User.is_active == True).order_by(User.name).all()
    form.owner_id.choices = [(str(u.id), f'{u.name} ({u.email})') for u in all_active_users]
    if form.validate_on_submit():
        owner_id = int(form.owner_id.data) if form.owner_id.data else user_id
        if user.role.name != 'Документовед':
            owner_id = user_id
        award = Award(
            user_id=owner_id,
            title=form.title.data,
            description=form.description.data,
            award_type=form.award_type.data,
            issuer=form.issuer.data,
            date_received=form.date_received.data,
            level=form.level.data,
        )
        db.session.add(award)
        db.session.commit()
        flash('Награда добавлена!', 'success')
        return redirect(url_for('rating.awards_list'))
    
    return render_template('rating/award_form.html', form=form, title='Добавить награду', user=user)

@rating_bp.route('/awards/<int:award_id>/edit', methods=['GET', 'POST'])
@jwt_required()
def edit_award(award_id):
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    award = Award.query.get_or_404(award_id)
    if user.role.name != 'Документовед' and award.user_id != user_id:
        flash('У вас нет доступа', 'danger')
        return redirect(url_for('rating.awards_list'))
    
    form = AwardForm()
    all_active_users = User.query.filter(User.is_active == True).order_by(User.name).all()
    form.owner_id.choices = [(str(u.id), f'{u.name} ({u.email})') for u in all_active_users]
    if form.validate_on_submit():
        award.title = form.title.data
        award.description = form.description.data
        award.award_type = form.award_type.data
        award.issuer = form.issuer.data
        award.date_received = form.date_received.data
        award.level = form.level.data
        if user.role.name == 'Документовед' and form.owner_id.data:
            award.user_id = int(form.owner_id.data)
        db.session.commit()
        flash('Награда обновлена!', 'success')
        return redirect(url_for('rating.awards_list'))
    
    elif request.method == 'GET':
        form.title.data = award.title
        form.description.data = award.description
        form.award_type.data = award.award_type
        form.issuer.data = award.issuer
        form.date_received.data = award.date_received
        form.level.data = award.level
        if user.role.name == 'Документовед':
            form.owner_id.data = str(award.user_id)
    
    return render_template('rating/award_form.html', form=form, title='Редактировать награду', award=award, user=user)

@rating_bp.route('/awards/<int:award_id>/delete', methods=['POST'])
@jwt_required()
def delete_award(award_id):
    user_id = get_jwt_identity()
    award = Award.query.get_or_404(award_id)
    if award.user_id != user_id:
        flash('У вас нет доступа', 'danger')
        return redirect(url_for('rating.awards_list'))
    
    db.session.delete(award)
    db.session.commit()
    flash('Награда удалена!', 'success')
    return redirect(url_for('rating.awards_list'))

# ==================== PUBLICATIONS ====================

@rating_bp.route('/publications')
@jwt_required()
def publications_list():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    search_form = SearchFilterForm()
    page = request.args.get('page', 1, type=int)

    users_list = User.query.filter(User.is_active == True).order_by(User.name).all()
    author_ids = request.args.getlist('author_ids', type=int)

    all_publications = []

    pub_type = request.args.get('pub_type', '')
    index_type = request.args.get('index_type', '')
    student_report = request.args.get('student_report', '')
    book_type = request.args.get('book_type', '')
    search_query = request.args.get('search_query', '')
    sort_by = request.args.get('sort_by', 'date_desc')
    date_range = request.args.get('date_range', 'all')
    date_from, date_to = parse_date_args(request.args.get('date_from'), request.args.get('date_to'))

    models_to_query = list(PUB_MODELS.values()) if not pub_type else [PUB_MODELS[pub_type]]

    pub_type_map = {v: k for k, v in PUB_MODELS.items()}

    for model in models_to_query:
        entity_type = pub_type_map[model]
        query = model.query.filter(filter_by_role(model, user, coauthor_entity_type=entity_type))
        if author_ids:
            query = query.filter(model.user_id.in_(author_ids))

        if index_type == 'scopus' and hasattr(model, 'scopus'):
            query = query.filter(model.scopus == True)
        elif index_type == 'vak' and hasattr(model, 'vak'):
            query = query.filter(model.vak == True)
        elif index_type == 'none' and hasattr(model, 'scopus') and hasattr(model, 'vak'):
            query = query.filter(model.scopus == False, model.vak == False)

        if student_report == 'yes':
            query = query.filter(model.student_report == True)
        elif student_report == 'no':
            query = query.filter(model.student_report == False)

        if book_type and hasattr(model, 'book_type'):
            query = query.filter(model.book_type == book_type)

        if search_query:
            author_fields = []
            if hasattr(model, 'authors'):
                author_fields.append(model.authors.ilike(f'%{search_query}%'))
            if hasattr(model, 'author_single'):
                author_fields.append(model.author_single.ilike(f'%{search_query}%'))
            if author_fields:
                query = query.filter(or_(model.title.ilike(f'%{search_query}%'), or_(*author_fields)))
            else:
                query = query.filter(model.title.ilike(f'%{search_query}%'))

        if date_range != 'all' and hasattr(model, 'publication_date'):
            if date_range == '3':
                cutoff = datetime.now() - timedelta(days=90)
                query = query.filter(model.publication_date >= cutoff.date())
            elif date_range == '6':
                cutoff = datetime.now() - timedelta(days=180)
                query = query.filter(model.publication_date >= cutoff.date())
            elif date_range == '12':
                cutoff = datetime.now() - timedelta(days=365)
                query = query.filter(model.publication_date >= cutoff.date())
            elif date_range == '36':
                cutoff = datetime.now() - timedelta(days=1095)
                query = query.filter(model.publication_date >= cutoff.date())
            elif date_range == 'academic_year':
                now = datetime.now()
                if now.month >= 9:
                    start = date(now.year, 9, 1)
                    end = date(now.year + 1, 8, 31)
                else:
                    start = date(now.year - 1, 9, 1)
                    end = date(now.year, 8, 31)
                query = query.filter(model.publication_date >= start, model.publication_date <= end)
            elif date_range == 'custom':
                if date_from:
                    query = query.filter(model.publication_date >= date_from)
                if date_to:
                    query = query.filter(model.publication_date <= date_to)

        all_publications.extend(query.all())

    if date_range != 'all':
        cutoff_date = None
        if date_range == '3':
            cutoff_date = datetime.now() - timedelta(days=90)
        elif date_range == '6':
            cutoff_date = datetime.now() - timedelta(days=180)
        elif date_range == '12':
            cutoff_date = datetime.now() - timedelta(days=365)
        elif date_range == '36':
            cutoff_date = datetime.now() - timedelta(days=1095)
        elif date_range == 'academic_year':
            now = datetime.now()
            if now.month >= 9:
                cutoff_date = date(now.year, 9, 1)
            else:
                cutoff_date = date(now.year - 1, 9, 1)

        filtered = []
        for pub in all_publications:
            if pub.publication_date is None:
                continue
            if cutoff_date and pub.publication_date < cutoff_date.date():
                continue
            if date_range == 'academic_year':
                now = datetime.now()
                if now.month >= 9:
                    end = date(now.year + 1, 8, 31)
                else:
                    end = date(now.year, 8, 31)
                if pub.publication_date > end:
                    continue
            if date_range == 'custom' and date_from and pub.publication_date < date_from:
                continue
            if date_range == 'custom' and date_to and pub.publication_date > date_to:
                continue
            filtered.append(pub)
        all_publications = filtered

    sort_column = request.args.get('sort_column', 'year')
    sort_dir = request.args.get('sort_dir', 'desc')
    type_names = {
        'publication_books': 'Книга',
        'publication_journal_articles': 'Журнал',
        'publication_collection_articles': 'Сборник научных статей',
        'publication_dissertations': 'Диссертация',
        'publication_abstracts': 'Автореферат',
        'publication_internets': 'Интернет',
    }

    if sort_column == 'year':
        all_publications.sort(key=lambda x: x.year or 0, reverse=(sort_dir == 'desc'))
    elif sort_column == 'title':
        all_publications.sort(key=lambda x: (x.title or '').lower(), reverse=(sort_dir == 'desc'))
    elif sort_column == 'type':
        all_publications.sort(key=lambda x: type_names.get(x.__tablename__, ''), reverse=(sort_dir == 'desc'))

    per_page = 10
    total = len(all_publications)
    start = (page - 1) * per_page
    end = start + per_page
    paginated = all_publications[start:end]

    class PaginatedResult:
        def __init__(self, items, total, page, per_page):
            self.items = items
            self.total = total
            self.page = page
            self.per_page = per_page
            self.pages = (total + per_page - 1) // per_page
            self.has_prev = page > 1
            self.has_next = page * per_page < total
            self.prev_num = page - 1
            self.next_num = page + 1

    publications = PaginatedResult(paginated, total, page, per_page)

    coauthors_map = {}
    supervisors_map = {}
    pub_type_lookup = {v: k for k, v in PUB_MODELS.items()}
    for pub in paginated:
        ptype = pub_type_lookup.get(type(pub))
        if ptype:
            coauthors = EntityCoauthor.query.filter_by(entity_type=ptype, entity_id=pub.id).all()
            if coauthors:
                coauthors_map[(type(pub).__tablename__, pub.id)] = [c.user for c in coauthors]
            supervisors = EntitySupervisor.query.filter_by(entity_type=ptype, entity_id=pub.id).all()
            if supervisors:
                supervisors_map[(type(pub).__tablename__, pub.id)] = [s.user for s in supervisors]

    return render_template(
        'rating/publications.html',
        publications=publications,
        search_form=search_form,
        coauthors_map=coauthors_map,
        supervisors_map=supervisors_map,
        users_list=users_list,
        author_ids=author_ids,
        student_report=student_report,
        book_type=book_type,
    )

@rating_bp.route('/publications/preview', methods=['POST'])
@jwt_required()
def preview_publication():
    data = request.get_json()
    pub_type = data.get('publication_type', '')
    model = PUB_MODELS.get(pub_type)
    if not model:
        return jsonify({'error': 'Unknown publication type'}), 400

    try:
        year_val = int(data.get('year', datetime.now().year)) if data.get('year') else datetime.now().year
        common = {
            'title': data.get('title', ''),
            'year': year_val,
            'isbn': data.get('isbn', ''),
            'student_report': data.get('student_report', False),
        }

        pub_date_str = data.get('publication_date', '')
        if pub_date_str:
            try:
                common['publication_date'] = datetime.strptime(pub_date_str, '%Y-%m-%d').date()
            except ValueError:
                common['publication_date'] = None
        else:
            common['publication_date'] = None

        if pub_type == 'book':
            temp_pub = model(
                **common,
                authors=data.get('authors', ''),
                edition=data.get('edition', ''),
                city=data.get('city', ''),
                publisher=data.get('publisher', ''),
                pages=data.get('pages', ''),
                book_type=data.get('book_type', ''),
            )
        elif pub_type == 'journal_article':
            temp_pub = model(
                **common,
                authors=data.get('authors', ''),
                journal_name=data.get('journal_name', ''),
                issue=data.get('issue', ''),
                pages=data.get('pages', ''),
                scopus=data.get('scopus', False),
                vak=data.get('vak', False),
            )
        elif pub_type == 'collection_article':
            temp_pub = model(
                **common,
                authors=data.get('authors', ''),
                collection_title=data.get('collection_title', ''),
                city=data.get('city', ''),
                publisher=data.get('publisher', ''),
                pages=data.get('pages', ''),
            )
        elif pub_type in ('dissertation', 'abstract'):
            temp_pub = model(
                **common,
                author_single=data.get('author_single', ''),
                degree=data.get('degree', ''),
                field=data.get('field', ''),
                specialty_code=data.get('specialty_code', ''),
                city=data.get('city', ''),
                pages=data.get('pages', ''),
            )
        elif pub_type == 'internet':
            temp_pub = model(
                **common,
                authors=data.get('authors', ''),
                site_name=data.get('site_name', ''),
                url=data.get('url', ''),
                access_date=data.get('access_date', ''),
            )
        else:
            return jsonify({'error': 'Unknown publication type'}), 400

        gost_string = generate_gost_string(temp_pub)
        return jsonify({'gost_string': gost_string})
    except Exception as e:
        return jsonify({'error': str(e)}), 400


@rating_bp.route('/publications/add', methods=['GET', 'POST'])
@jwt_required()
def add_publication():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    form = PublicationForm()
    users = User.query.join(User.role).filter(User.is_active == True, User.id != user_id, db.text("roles.name != 'Руководитель'")).order_by(User.name).all()
    form.coauthor_ids.choices = [(str(u.id), f'{u.name} ({u.email})') for u in users]

    all_active_users = User.query.filter(User.is_active == True).order_by(User.name).all()
    form.owner_id.choices = [(str(u.id), f'{u.name} ({u.email})') for u in all_active_users]
    form.supervisor_id.choices = [('', '---')] + [(str(u.id), f'{u.name} ({u.email})') for u in all_active_users]

    if form.validate_on_submit():
        pub_type = form.publication_type.data
        model = PUB_MODELS.get(pub_type)
        if not model:
            flash('Неизвестный тип публикации', 'danger')
            return redirect(url_for('rating.publications_list'))

        pub_date = form.publication_date.data if form.publication_date.data else None

        owner_id = int(form.owner_id.data) if form.owner_id.data else user_id
        if user.role.name != 'Документовед':
            owner_id = user_id

        common = {
            'user_id': owner_id,
            'title': form.title.data,
            'year': form.year.data,
            'publication_date': pub_date,
            'isbn': form.isbn.data,
            'student_report': form.student_report.data,
        }

        if pub_type == 'book':
            pub = model(**common, authors=form.authors.data, edition=form.edition.data, city=form.city.data, publisher=form.publisher.data, pages=form.pages.data, book_type=form.book_type.data)
        elif pub_type == 'journal_article':
            pub = model(**common, authors=form.authors.data, journal_name=form.journal_name.data, issue=form.issue.data, pages=form.pages.data, scopus=form.scopus.data, vak=form.vak.data)
        elif pub_type == 'collection_article':
            pub = model(**common, authors=form.authors.data, collection_title=form.collection_title.data, city=form.city.data, publisher=form.publisher.data, pages=form.pages.data)
        elif pub_type == 'dissertation':
            pub = model(**common, author_single=form.author_single.data, degree=form.degree.data, field=form.field.data, specialty_code=form.specialty_code.data, city=form.city.data, pages=form.pages.data)
        elif pub_type == 'abstract':
            pub = model(**common, author_single=form.author_single.data, degree=form.degree.data, field=form.field.data, specialty_code=form.specialty_code.data, city=form.city.data, pages=form.pages.data)
        elif pub_type == 'internet':
            pub = model(**common, authors=form.authors.data, site_name=form.site_name.data, url=form.url.data, access_date=form.access_date.data)
        else:
            flash('Неизвестный тип публикации', 'danger')
            return redirect(url_for('rating.publications_list'))

        pub.gost_string = generate_gost_string(pub)
        db.session.add(pub)
        db.session.flush()

        coauthor_ids = form.coauthor_ids.data or []
        for coauthor_id in coauthor_ids:
            coauthor_id = int(coauthor_id)
            if coauthor_id != user_id:
                db.session.add(EntityCoauthor(
                    entity_type=pub_type,
                    entity_id=pub.id,
                    user_id=coauthor_id,
                ))
        if form.supervisor_id.data:
            db.session.add(EntitySupervisor(
                entity_type=pub_type,
                entity_id=pub.id,
                user_id=int(form.supervisor_id.data),
            ))
        db.session.commit()
        flash('Публикация добавлена!', 'success')
        return redirect(url_for('rating.publications_list'))

    return render_template('rating/publication_form.html', form=form, title='Добавить публикацию', users=users, user=user)


@rating_bp.route('/publications/<int:pub_id>/edit', methods=['GET', 'POST'])
@jwt_required()
def edit_publication(pub_id):
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    pub_type = request.args.get('pub_type', '')
    pub_record = _get_publication_by_id(pub_id, pub_type)
    if not pub_record:
        flash('Публикация не найдена', 'danger')
        return redirect(url_for('rating.publications_list'))
    if user.role.name != 'Документовед' and pub_record.user_id != user_id:
        flash('У вас нет доступа', 'danger')
        return redirect(url_for('rating.publications_list'))

    for key, model in PUB_MODELS.items():
        if isinstance(pub_record, model):
            form_pub_type = key
            break
    else:
        form_pub_type = 'book'

    form = PublicationForm()
    users_list = User.query.join(User.role).filter(User.is_active == True, User.id != user_id, db.text("roles.name != 'Руководитель'")).order_by(User.name).all()
    form.coauthor_ids.choices = [(str(u.id), f'{u.name} ({u.email})') for u in users_list]

    all_active_users = User.query.filter(User.is_active == True).order_by(User.name).all()
    form.owner_id.choices = [(str(u.id), f'{u.name} ({u.email})') for u in all_active_users]
    form.supervisor_id.choices = [('', '---')] + [(str(u.id), f'{u.name} ({u.email})') for u in all_active_users]

    if form.validate_on_submit():
        pub_date = form.publication_date.data if form.publication_date.data else None

        pub_record.title = form.title.data
        pub_record.year = form.year.data
        pub_record.publication_date = pub_date
        pub_record.isbn = form.isbn.data
        pub_record.student_report = form.student_report.data

        if user.role.name == 'Документовед' and form.owner_id.data:
            pub_record.user_id = int(form.owner_id.data)

        if hasattr(pub_record, 'authors'):
            pub_record.authors = form.authors.data
        if hasattr(pub_record, 'author_single'):
            pub_record.author_single = form.author_single.data
        if hasattr(pub_record, 'edition'):
            pub_record.edition = form.edition.data
        if hasattr(pub_record, 'city'):
            pub_record.city = form.city.data
        if hasattr(pub_record, 'publisher'):
            pub_record.publisher = form.publisher.data
        if hasattr(pub_record, 'pages'):
            pub_record.pages = form.pages.data
        if hasattr(pub_record, 'journal_name'):
            pub_record.journal_name = form.journal_name.data
        if hasattr(pub_record, 'issue'):
            pub_record.issue = form.issue.data
        if hasattr(pub_record, 'scopus'):
            pub_record.scopus = form.scopus.data
        if hasattr(pub_record, 'vak'):
            pub_record.vak = form.vak.data
        if hasattr(pub_record, 'collection_title'):
            pub_record.collection_title = form.collection_title.data
        if hasattr(pub_record, 'degree'):
            pub_record.degree = form.degree.data
        if hasattr(pub_record, 'field'):
            pub_record.field = form.field.data
        if hasattr(pub_record, 'specialty_code'):
            pub_record.specialty_code = form.specialty_code.data
        if hasattr(pub_record, 'site_name'):
            pub_record.site_name = form.site_name.data
        if hasattr(pub_record, 'url'):
            pub_record.url = form.url.data
        if hasattr(pub_record, 'access_date'):
            pub_record.access_date = form.access_date.data
        if hasattr(pub_record, 'book_type'):
            pub_record.book_type = form.book_type.data

        pub_record.gost_string = generate_gost_string(pub_record)

        EntityCoauthor.query.filter_by(entity_type=form_pub_type, entity_id=pub_record.id).delete()
        EntitySupervisor.query.filter_by(entity_type=form_pub_type, entity_id=pub_record.id).delete()
        coauthor_ids = form.coauthor_ids.data or []
        for coauthor_id in coauthor_ids:
            coauthor_id = int(coauthor_id)
            if coauthor_id != user_id:
                db.session.add(EntityCoauthor(
                    entity_type=form_pub_type,
                    entity_id=pub_record.id,
                    user_id=coauthor_id,
                ))
        if form.supervisor_id.data:
            db.session.add(EntitySupervisor(
                entity_type=form_pub_type,
                entity_id=pub_record.id,
                user_id=int(form.supervisor_id.data),
            ))
        db.session.commit()
        flash('Публикация обновлена!', 'success')
        return redirect(url_for('rating.publications_list'))

    if request.method == 'GET':
        form.publication_type.data = form_pub_type
        form.title.data = pub_record.title
        form.year.data = pub_record.year
        form.publication_date.data = pub_record.publication_date
        form.isbn.data = pub_record.isbn
        form.student_report.data = pub_record.student_report
        if user.role.name == 'Документовед':
            form.owner_id.data = str(pub_record.user_id)
        existing_supervisors = EntitySupervisor.query.filter_by(entity_type=form_pub_type, entity_id=pub_record.id).all()
        if existing_supervisors:
            form.supervisor_id.data = str(existing_supervisors[0].user_id)
        existing_coauthors = EntityCoauthor.query.filter_by(entity_type=form_pub_type, entity_id=pub_record.id).all()
        form.coauthor_ids.data = [str(c.user_id) for c in existing_coauthors]
        if hasattr(pub_record, 'authors'):
            form.authors.data = pub_record.authors
        if hasattr(pub_record, 'author_single'):
            form.author_single.data = pub_record.author_single
        if hasattr(pub_record, 'edition'):
            form.edition.data = pub_record.edition
        if hasattr(pub_record, 'city'):
            form.city.data = pub_record.city
        if hasattr(pub_record, 'publisher'):
            form.publisher.data = pub_record.publisher
        if hasattr(pub_record, 'pages'):
            form.pages.data = pub_record.pages
        if hasattr(pub_record, 'journal_name'):
            form.journal_name.data = pub_record.journal_name
        if hasattr(pub_record, 'issue'):
            form.issue.data = pub_record.issue
        if hasattr(pub_record, 'scopus'):
            form.scopus.data = pub_record.scopus
        if hasattr(pub_record, 'vak'):
            form.vak.data = pub_record.vak
        if hasattr(pub_record, 'collection_title'):
            form.collection_title.data = pub_record.collection_title
        if hasattr(pub_record, 'degree'):
            form.degree.data = pub_record.degree
        if hasattr(pub_record, 'field'):
            form.field.data = pub_record.field
        if hasattr(pub_record, 'specialty_code'):
            form.specialty_code.data = pub_record.specialty_code
        if hasattr(pub_record, 'site_name'):
            form.site_name.data = pub_record.site_name
        if hasattr(pub_record, 'url'):
            form.url.data = pub_record.url
        if hasattr(pub_record, 'access_date'):
            form.access_date.data = pub_record.access_date
        if hasattr(pub_record, 'book_type'):
            form.book_type.data = pub_record.book_type

    return render_template('rating/publication_form.html', form=form, title='Редактировать публикацию', publication=pub_record, pub_type_label=PUB_TYPE_LABELS.get(form_pub_type, form_pub_type), users=users_list, user=user)


def _get_publication_by_id(pub_id, pub_type=None):
    if pub_type and pub_type in PUB_MODELS:
        return PUB_MODELS[pub_type].query.get(pub_id)
    for model in PUB_MODELS.values():
        pub = model.query.get(pub_id)
        if pub:
            return pub
    return None


@rating_bp.route('/publications/<int:pub_id>', methods=['GET'])
@jwt_required()
def view_publication(pub_id):
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    pub_type = request.args.get('pub_type', '')
    publication = _get_publication_by_id(pub_id, pub_type)
    if not publication:
        flash('Публикация не найдена', 'danger')
        return redirect(url_for('rating.publications_list'))

    is_owner = publication.user_id == user_id
    if user.role.name != 'Документовед' and not is_owner:
        coauthor_ids = [c.user_id for c in EntityCoauthor.query.filter_by(
            entity_type=pub_type, entity_id=publication.id
        ).all()]
        if user_id not in coauthor_ids:
            flash('У вас нет доступа', 'danger')
            return redirect(url_for('rating.publications_list'))

    pub_type_key = None
    for key, model in PUB_MODELS.items():
        if isinstance(publication, model):
            pub_type_key = key
            break

    coauthors = EntityCoauthor.query.filter_by(entity_type=pub_type_key, entity_id=publication.id).all()
    return render_template('rating/publication_view.html', publication=publication, pub_type_key=pub_type_key, coauthors=coauthors)


@rating_bp.route('/publications/<int:pub_id>/delete', methods=['POST'])
@jwt_required()
def delete_publication(pub_id):
    user_id = get_jwt_identity()
    pub_type = request.args.get('pub_type', '')
    publication = _get_publication_by_id(pub_id, pub_type)
    if not publication:
        flash('Публикация не найдена', 'danger')
        return redirect(url_for('rating.publications_list'))
    if publication.user_id != user_id:
        flash('У вас нет доступа', 'danger')
        return redirect(url_for('rating.publications_list'))

    db.session.delete(publication)
    db.session.commit()
    flash('Публикация удалена!', 'success')
    return redirect(url_for('rating.publications_list'))


# ==================== CONFERENCES ====================

@rating_bp.route('/conferences')
@jwt_required()
def conferences_list():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    search_form = SearchFilterForm()
    page = request.args.get('page', 1, type=int)

    users_list = User.query.filter(User.is_active == True).order_by(User.name).all()

    query = Conference.query.filter(filter_by_role(Conference, user)).filter(Conference.status == 'active')

    author_ids = request.args.getlist('author_ids', type=int)
    if author_ids:
        query = query.filter(Conference.user_id.in_(author_ids))

    search_query = request.args.get('search_query', '')
    if search_query:
        query = query.filter(or_(
            Conference.name.ilike(f'%{search_query}%'),
            Conference.paper_title.ilike(f'%{search_query}%')
        ))

    date_range = request.args.get('date_range', 'all')
    date_from, date_to = parse_date_args(request.args.get('date_from'), request.args.get('date_to'))
    query = apply_date_filter(query, Conference.conference_date, date_range, date_from, date_to)

    sort_column = request.args.get('sort_column', 'conference_date')
    sort_dir = request.args.get('sort_dir', 'desc')
    col_map = {
        'name': Conference.name,
        'role': Conference.role,
        'paper_title': Conference.paper_title,
        'conference_date': Conference.conference_date,
        'coauthors': Conference.coauthors,
    }
    col = col_map.get(sort_column, Conference.conference_date)
    query = query.order_by(col.desc() if sort_dir == 'desc' else col.asc())

    conferences = query.paginate(page=page, per_page=10)

    return render_template('rating/conferences.html', conferences=conferences, search_form=search_form, users_list=users_list, author_ids=author_ids)

@rating_bp.route('/conferences/add', methods=['GET', 'POST'])
@jwt_required()
def add_conference():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    form = ConferenceForm()
    all_active_users = User.query.filter(User.is_active == True).order_by(User.name).all()
    form.owner_id.choices = [(str(u.id), f'{u.name} ({u.email})') for u in all_active_users]
    if form.validate_on_submit():
        owner_id = int(form.owner_id.data) if form.owner_id.data else user_id
        if user.role.name != 'Документовед':
            owner_id = user_id
        conference = Conference(
            user_id=owner_id,
            name=form.name.data,
            role=form.role.data,
            paper_title=form.paper_title.data,
            conference_date=form.conference_date.data,
            location=form.location.data,
            conference_url=form.conference_url.data,
            description=form.description.data,
            coauthors=form.coauthors.data
        )
        db.session.add(conference)
        db.session.commit()
        flash('Конференция добавлена!', 'success')
        return redirect(url_for('rating.conferences_list'))
    
    return render_template('rating/conference_form.html', form=form, title='Добавить конференцию', user=user)

@rating_bp.route('/conferences/<int:conf_id>/edit', methods=['GET', 'POST'])
@jwt_required()
def edit_conference(conf_id):
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    conference = Conference.query.get_or_404(conf_id)
    if user.role.name != 'Документовед' and conference.user_id != user_id:
        flash('У вас нет доступа', 'danger')
        return redirect(url_for('rating.conferences_list'))
    
    form = ConferenceForm()
    all_active_users = User.query.filter(User.is_active == True).order_by(User.name).all()
    form.owner_id.choices = [(str(u.id), f'{u.name} ({u.email})') for u in all_active_users]
    if form.validate_on_submit():
        conference.name = form.name.data
        conference.role = form.role.data
        conference.paper_title = form.paper_title.data
        conference.conference_date = form.conference_date.data
        conference.location = form.location.data
        conference.conference_url = form.conference_url.data
        conference.description = form.description.data
        conference.coauthors=form.coauthors.data
        if user.role.name == 'Документовед' and form.owner_id.data:
            conference.user_id = int(form.owner_id.data)
        db.session.commit()
        flash('Конференция обновлена!', 'success')
        return redirect(url_for('rating.conferences_list'))
    
    elif request.method == 'GET':
        form.name.data = conference.name
        form.role.data = conference.role
        form.paper_title.data = conference.paper_title
        form.conference_date.data = conference.conference_date
        form.location.data = conference.location
        form.conference_url.data = conference.conference_url
        form.description.data = conference.description
        form.coauthors.data=conference.coauthors
        if user.role.name == 'Документовед':
            form.owner_id.data = str(conference.user_id)
    
    return render_template('rating/conference_form.html', form=form, title='Редактировать конференцию', conference=conference, user=user)

@rating_bp.route('/conferences/<int:conf_id>/delete', methods=['POST'])
@jwt_required()
def delete_conference(conf_id):
    user_id = get_jwt_identity()
    conference = Conference.query.get_or_404(conf_id)
    if conference.user_id != user_id:
        flash('У вас нет доступа', 'danger')
        return redirect(url_for('rating.conferences_list'))
    
    db.session.delete(conference)
    db.session.commit()
    flash('Конференция удалена!', 'success')
    return redirect(url_for('rating.conferences_list'))


# ==================== TRAININGS ====================

@rating_bp.route('/trainings')
@jwt_required()
def trainings_list():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    search_form = SearchFilterForm()
    page = request.args.get('page', 1, type=int)

    users_list = User.query.filter(User.is_active == True).order_by(User.name).all()

    query = Training.query.filter(filter_by_role(Training, user)).filter(Training.status == 'active')

    state_filter = request.args.get('state_issued', '')
    if state_filter == 'yes':
        query = query.filter(Training.state_issued == True)
    elif state_filter == 'no':
        query = query.filter(Training.state_issued == False)

    author_ids = request.args.getlist('author_ids', type=int)
    if author_ids:
        query = query.filter(Training.user_id.in_(author_ids))

    search_query = request.args.get('search_query', '')
    if search_query:
        query = query.filter(or_(
            Training.title.ilike(f'%{search_query}%'),
            Training.organization.ilike(f'%{search_query}%')
        ))

    date_range = request.args.get('date_range', 'all')
    date_from, date_to = parse_date_args(request.args.get('date_from'), request.args.get('date_to'))
    query = apply_date_filter(query, Training.end_date, date_range, date_from, date_to)

    sort_column = request.args.get('sort_column', 'end_date')
    sort_dir = request.args.get('sort_dir', 'desc')
    col_map = {
        'title': Training.title,
        'organization': Training.organization,
        'end_date': Training.end_date,
        'city': Training.city,
        'level': Training.level,
    }
    col = col_map.get(sort_column, Training.end_date)
    query = query.order_by(col.desc() if sort_dir == 'desc' else col.asc())

    trainings = query.paginate(page=page, per_page=10)

    return render_template('rating/trainings.html', trainings=trainings, search_form=search_form, users_list=users_list, author_ids=author_ids, state_filter=state_filter)

@rating_bp.route('/trainings/add', methods=['GET', 'POST'])
@jwt_required()
def add_training():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    form = TrainingForm()
    all_active_users = User.query.filter(User.is_active == True).order_by(User.name).all()
    form.owner_id.choices = [(str(u.id), f'{u.name} ({u.email})') for u in all_active_users]
    if form.validate_on_submit():
        owner_id = int(form.owner_id.data) if form.owner_id.data else user_id
        if user.role.name != 'Документовед':
            owner_id = user_id
        training = Training(
            user_id=owner_id,
            title=form.title.data,
            organization=form.organization.data,
            city=form.city.data,
            training_type=form.training_type.data,
            start_date=form.start_date.data,
            end_date=form.end_date.data,
            duration_hours=form.duration_hours.data,
            certificate_number=form.certificate_number.data,
            certificate_url=form.certificate_url.data,
            description=form.description.data,
            level=form.level.data,
            state_issued=form.state_issued.data
        )
        db.session.add(training)
        db.session.commit()
        flash('Повышение квалификации добавлено!', 'success')
        return redirect(url_for('rating.trainings_list'))

    return render_template('rating/training_form.html', form=form, title='Добавить повышение квалификации', user=user)

@rating_bp.route('/trainings/<int:training_id>/edit', methods=['GET', 'POST'])
@jwt_required()
def edit_training(training_id):
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    training = Training.query.get_or_404(training_id)
    if user.role.name != 'Документовед' and training.user_id != user_id:
        flash('У вас нет доступа', 'danger')
        return redirect(url_for('rating.trainings_list'))
    
    form = TrainingForm()
    all_active_users = User.query.filter(User.is_active == True).order_by(User.name).all()
    form.owner_id.choices = [(str(u.id), f'{u.name} ({u.email})') for u in all_active_users]
    if form.validate_on_submit():
        training.title = form.title.data
        training.organization = form.organization.data
        training.city = form.city.data
        training.training_type = form.training_type.data
        training.start_date = form.start_date.data
        training.end_date = form.end_date.data
        training.duration_hours = form.duration_hours.data
        training.certificate_number = form.certificate_number.data
        training.certificate_url = form.certificate_url.data
        training.description = form.description.data
        training.level = form.level.data
        training.state_issued = form.state_issued.data
        if user.role.name == 'Документовед' and form.owner_id.data:
            training.user_id = int(form.owner_id.data)
        db.session.commit()
        flash('Повышение квалификации обновлено!', 'success')
        return redirect(url_for('rating.trainings_list'))
    
    elif request.method == 'GET':
        form.title.data = training.title
        form.organization.data = training.organization
        form.city.data = training.city
        form.training_type.data = training.training_type
        form.start_date.data = training.start_date
        form.end_date.data = training.end_date
        form.duration_hours.data = training.duration_hours
        form.certificate_number.data = training.certificate_number
        form.certificate_url.data = training.certificate_url
        form.description.data = training.description
        form.level.data = training.level
        form.state_issued.data = training.state_issued
        if user.role.name == 'Документовед':
            form.owner_id.data = str(training.user_id)
    
    return render_template('rating/training_form.html', form=form, title='Редактировать повышение квалификации', training=training, user=user)

@rating_bp.route('/trainings/<int:training_id>/delete', methods=['POST'])
@jwt_required()
def delete_training(training_id):
    user_id = get_jwt_identity()
    training = Training.query.get_or_404(training_id)
    if training.user_id != user_id:
        flash('У вас нет доступа', 'danger')
        return redirect(url_for('rating.trainings_list'))
    
    db.session.delete(training)
    db.session.commit()
    flash('Повышение квалификации удалено!', 'success')
    return redirect(url_for('rating.trainings_list'))

# ==================== TEMPLATES (AJAX + Management) ====================

ENTITY_TYPE_CHOICES = {
    'publication': 'Публикация',
    'award': 'Награда',
    'conference': 'Конференция',
    'training': 'Повышение квалификации',
}

FIELD_LABELS = {
    'publication': {
        'publication_type': 'Тип публикации',
        'title': 'Название',
        'year': 'Год',
        'publication_date': 'Дата публикации',
        'isbn': 'ISBN',
        'authors': 'Авторы',
        'author_single': 'Автор',
        'edition': 'Номер издания',
        'city': 'Город издания',
        'publisher': 'Издательство',
        'pages': 'Страницы',
        'journal_name': 'Название журнала',
        'issue': 'Номер выпуска',
        'collection_title': 'Название сборника',
        'degree': 'Учёная степень',
        'field': 'Отрасль наук',
        'specialty_code': 'Код специальности',
        'site_name': 'Название сайта',
        'url': 'Гиперссылка',
        'access_date': 'Дата обращения',
        'book_type': 'Тип книги',
    },
    'award': {
        'title': 'Название награды',
        'description': 'Описание',
        'award_type': 'Тип награды',
        'issuer': 'Выдавшая организация',
        'date_received': 'Дата получения',
        'points': 'Очки рейтинга',
        'level': 'Уровень награды',
    },
    'conference': {
        'name': 'Название конференции',
        'role': 'Роль',
        'paper_title': 'Название доклада/статьи',
        'conference_date': 'Дата конференции',
        'location': 'Место проведения',
        'conference_url': 'Сайт конференции',
        'description': 'Описание',
        'points': 'Очки рейтинга',
        'coauthors': 'Соавторы',
    },
    'training': {
        'title': 'Название курса/тренинга',
        'organization': 'Организация',
        'city': 'Город проведения',
        'training_type': 'Тип обучения',
        'start_date': 'Дата начала',
        'end_date': 'Дата окончания',
        'duration_hours': 'Продолжительность (часов)',
        'certificate_number': 'Номер сертификата',
        'certificate_url': 'Ссылка на сертификат',
        'description': 'Описание',
        'points': 'Очки рейтинга',
        'level': 'Уровень',
        'state_issued': 'Государственного образца',
    },
}

PUBLICATION_TYPE_LABELS = dict(PUBLICATION_TYPES)
AWARD_TYPE_LABELS = dict(AWARD_TYPES)
CONFERENCE_ROLE_LABELS = dict(CONFERENCE_ROLES)
TRAINING_TYPE_LABELS = dict([
    ('course', 'Курс'),
    ('workshop', 'Мастер-класс'),
    ('seminar', 'Семинар'),
    ('webinar', 'Вебинар'),
    ('certification', 'Сертификация'),
])
LEVEL_LABELS = dict(AWARD_LEVELS)

@rating_bp.route('/publications/api/templates', methods=['GET'])
@jwt_required()
def api_list_templates():
    entity_type = request.args.get('entity_type', '')
    sub_type = request.args.get('sub_type', '')
    query = RatingTemplate.query
    if entity_type:
        query = query.filter(RatingTemplate.entity_type == entity_type)
    if sub_type:
        query = query.filter(RatingTemplate.sub_type == sub_type)
    templates = query.order_by(RatingTemplate.name).all()
    return jsonify([{
        'id': t.id,
        'name': t.name,
        'entity_type': t.entity_type,
        'sub_type': t.sub_type,
        'created_at': t.created_at.strftime('%d.%m.%Y %H:%M'),
    } for t in templates])


@rating_bp.route('/publications/api/templates', methods=['POST'])
@jwt_required()
@roles_required('Документовед')
def api_save_template():
    data = request.get_json()
    if not data or not data.get('name') or not data.get('entity_type'):
        return jsonify({'error': 'Не указано имя или тип шаблона'}), 400
    template = RatingTemplate(
        name=data['name'],
        entity_type=data['entity_type'],
        sub_type=data.get('sub_type', ''),
        template_data=data.get('template_data', {}),
    )
    db.session.add(template)
    db.session.commit()
    return jsonify({'id': template.id, 'name': template.name})


@rating_bp.route('/publications/api/templates/<int:template_id>', methods=['GET'])
@jwt_required()
def api_get_template(template_id):
    template = RatingTemplate.query.get_or_404(template_id)
    return jsonify({
        'id': template.id,
        'name': template.name,
        'entity_type': template.entity_type,
        'sub_type': template.sub_type,
        'template_data': template.template_data,
    })


@rating_bp.route('/publications/api/templates/<int:template_id>', methods=['PUT'])
@jwt_required()
@roles_required('Документовед')
def api_update_template(template_id):
    template = RatingTemplate.query.get_or_404(template_id)
    data = request.get_json()
    if data.get('name'):
        template.name = data['name']
    if data.get('template_data'):
        template.template_data = data['template_data']
    if data.get('sub_type') is not None:
        template.sub_type = data['sub_type']
    db.session.commit()
    return jsonify({'ok': True})


@rating_bp.route('/publications/api/templates/<int:template_id>', methods=['DELETE', 'POST'])
@jwt_required()
@roles_required('Документовед')
def api_delete_template(template_id):
    if request.method == 'POST' and request.form.get('_method') != 'DELETE':
        if not request.is_json:
            return jsonify({'error': 'Method not allowed'}), 405
    template = RatingTemplate.query.get_or_404(template_id)
    db.session.delete(template)
    db.session.commit()
    if request.is_json or request.method == 'DELETE':
        return jsonify({'ok': True})
    flash('Шаблон удалён', 'success')
    return redirect(url_for('rating.templates_list'))


@rating_bp.route('/templates/create', methods=['GET', 'POST'])
@jwt_required()
@roles_required('Документовед')
def create_template():
    entity_type = request.args.get('entity_type', 'award')
    sub_type = request.args.get('sub_type', '')

    form_map = {
        'publication': (PublicationForm, 'rating/publication_form.html'),
        'award': (AwardForm, 'rating/award_form.html'),
        'conference': (ConferenceForm, 'rating/conference_form.html'),
        'training': (TrainingForm, 'rating/training_form.html'),
    }

    if entity_type not in form_map:
        flash('Неизвестный тип', 'danger')
        return redirect(url_for('rating.templates_list'))

    form_class, tmpl = form_map[entity_type]
    form = form_class()

    if request.method == 'POST':
        for field in form:
            field.validators = [Optional()]
        if form.validate_on_submit():
            template_data = {}
            for field in form:
                fname = field.name
                if fname in ('submit', 'csrf_token'):
                    continue
                if fname == 'publication_type':
                    sub_type = field.data or ''
                    continue
                val = field.data
                if val is not None:
                    if isinstance(val, date):
                        val = val.isoformat()
                    template_data[fname] = val

            template = RatingTemplate(
                name=request.form.get('template_name', 'Новый шаблон'),
                entity_type=entity_type,
                sub_type=sub_type if entity_type == 'publication' else '',
                template_data=template_data,
            )
            db.session.add(template)
            db.session.commit()
            flash('Шаблон «{}» сохранён!'.format(template.name), 'success')
            return redirect(url_for('rating.templates_list'))

    users = User.query.join(User.role).filter(User.is_active == True, db.text("roles.name != 'Руководитель'")).order_by(User.name).all()
    current_user_id = get_jwt_identity()
    return render_template(tmpl,
        form=form,
        title='Создать шаблон — {}'.format(ENTITY_TYPE_CHOICES.get(entity_type, entity_type)),
        template_mode=True,
        cancel_url=url_for('rating.templates_list'),
        entity_type=entity_type,
        users=users,
        current_user_id=current_user_id)


@rating_bp.route('/templates')
@jwt_required()
def templates_list():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    templates = RatingTemplate.query.order_by(RatingTemplate.updated_at.desc()).all()
    return render_template('rating/templates.html', templates=templates, entity_labels=ENTITY_TYPE_CHOICES, role=user.role, PUBLICATION_TYPE_LABELS=PUBLICATION_TYPE_LABELS)


@rating_bp.route('/templates/<int:template_id>')
@jwt_required()
def template_view(template_id):
    template = RatingTemplate.query.get_or_404(template_id)
    labels = FIELD_LABELS.get(template.entity_type, {})

    choice_labels = {}
    if template.entity_type == 'publication':
        choice_labels = {
            'publication_type': PUBLICATION_TYPE_LABELS,
            'degree': dict(DEGREES),
            'field': dict(FIELDS_OF_STUDY),
        }
    elif template.entity_type == 'award':
        choice_labels = {
            'award_type': AWARD_TYPE_LABELS,
            'level': LEVEL_LABELS,
        }
    elif template.entity_type == 'conference':
        choice_labels = {
            'role': CONFERENCE_ROLE_LABELS,
        }
    elif template.entity_type == 'training':
        choice_labels = {
            'training_type': TRAINING_TYPE_LABELS,
            'level': LEVEL_LABELS,
        }

    return render_template('rating/template_view.html',
        template=template,
        labels=labels,
        choice_labels=choice_labels,
        entity_label=ENTITY_TYPE_CHOICES.get(template.entity_type, template.entity_type),
        PUBLICATION_TYPE_LABELS=PUBLICATION_TYPE_LABELS,
    )


"""Экспорт в Excel"""

@rating_bp.route('/awards/export/excel')
@jwt_required()
def export_awards_excel():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    query = Award.query.filter(filter_by_role(Award, user)).filter(Award.status == 'active')

    author_ids = request.args.getlist('author_ids', type=int)
    if author_ids:
        query = query.filter(Award.user_id.in_(author_ids))

    search_query = request.args.get('search_query', '')
    if search_query:
        query = query.filter(or_(
            Award.title.ilike(f'%{search_query}%'),
            Award.description.ilike(f'%{search_query}%')
        ))

    date_range = request.args.get('date_range', 'all')
    date_from, date_to = parse_date_args(request.args.get('date_from'), request.args.get('date_to'))
    query = apply_date_filter(query, Award.date_received, date_range, date_from, date_to)

    awards = query.all()
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Награды"
    
    headers = ["ID", "Название", "Тип", "Организация", "Дата получения", "Уровень"]
    ws.append(headers)
    
    header_fill = PatternFill(start_color="007BFF", end_color="007BFF", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
    
    for award in awards:
        ws.append([
            award.id,
            award.title,
            award.award_type,
            award.issuer or "",
            award.date_received.strftime("%d.%m.%Y"),
            award.level or ""
        ])
    
    ws.column_dimensions['A'].width = 5
    ws.column_dimensions['B'].width = 30
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 25
    ws.column_dimensions['E'].width = 15
    ws.column_dimensions['F'].width = 10
    
    bytes_stream = BytesIO()
    wb.save(bytes_stream)
    bytes_stream.seek(0)
    
    return send_file(
        bytes_stream,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f"awards_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    )


@rating_bp.route('/publications/export/excel', methods=['GET'])
@jwt_required()
def export_publications_excel():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    pub_type = request.args.get('pub_type', '')
    index_type = request.args.get('index_type', '')
    search_query = request.args.get('search_query', '')
    date_range = request.args.get('date_range', 'all')
    date_from, date_to = parse_date_args(request.args.get('date_from'), request.args.get('date_to'))
    author_ids = request.args.getlist('author_ids', type=int)
    pub_type_map_rev = {v: k for k, v in PUB_MODELS.items()}

    models_to_query = list(PUB_MODELS.values()) if not pub_type else [PUB_MODELS[pub_type]]
    all_publications = []

    for model in models_to_query:
        entity_type = pub_type_map_rev[model]
        query = model.query.filter(filter_by_role(model, user, coauthor_entity_type=entity_type))
        if author_ids:
            query = query.filter(model.user_id.in_(author_ids))

        if index_type == 'scopus' and hasattr(model, 'scopus'):
            query = query.filter(model.scopus == True)
        elif index_type == 'vak' and hasattr(model, 'vak'):
            query = query.filter(model.vak == True)
        elif index_type == 'none' and hasattr(model, 'scopus') and hasattr(model, 'vak'):
            query = query.filter(model.scopus == False, model.vak == False)

        if search_query:
            author_fields = []
            if hasattr(model, 'authors'):
                author_fields.append(model.authors.ilike(f'%{search_query}%'))
            if hasattr(model, 'author_single'):
                author_fields.append(model.author_single.ilike(f'%{search_query}%'))
            if author_fields:
                query = query.filter(or_(model.title.ilike(f'%{search_query}%'), or_(*author_fields)))
            else:
                query = query.filter(model.title.ilike(f'%{search_query}%'))

        if date_range != 'all' and hasattr(model, 'publication_date'):
            if date_range == '3':
                cutoff = datetime.now() - timedelta(days=90)
                query = query.filter(model.publication_date >= cutoff.date())
            elif date_range == '6':
                cutoff = datetime.now() - timedelta(days=180)
                query = query.filter(model.publication_date >= cutoff.date())
            elif date_range == '12':
                cutoff = datetime.now() - timedelta(days=365)
                query = query.filter(model.publication_date >= cutoff.date())
            elif date_range == '36':
                cutoff = datetime.now() - timedelta(days=1095)
                query = query.filter(model.publication_date >= cutoff.date())
            elif date_range == '60':
                cutoff = datetime.now() - timedelta(days=1825)
                query = query.filter(model.publication_date >= cutoff.date())
            elif date_range == 'academic_year':
                now = datetime.now()
                if now.month >= 9:
                    start = date(now.year, 9, 1)
                    end = date(now.year + 1, 8, 31)
                else:
                    start = date(now.year - 1, 9, 1)
                    end = date(now.year, 8, 31)
                query = query.filter(model.publication_date >= start, model.publication_date <= end)
            elif date_range == 'custom':
                if date_from:
                    query = query.filter(model.publication_date >= date_from)
                if date_to:
                    query = query.filter(model.publication_date <= date_to)

        all_publications.extend(query.all())

    wb = Workbook()
    ws = wb.active
    ws.title = 'Публикации'

    headers = ['№', 'Тип', 'Название', 'ГОСТ-строка']
    ws.append(headers)

    header_fill = PatternFill(start_color='28A745', end_color='28A745', fill_type='solid')
    header_font = Font(bold=True, color='FFFFFF')

    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')

    for idx, pub in enumerate(all_publications, 1):
        pub_type_label = ''
        for key, model in PUB_MODELS.items():
            if isinstance(pub, model):
                pub_type_label = PUB_TYPE_LABELS.get(key, key)
                break
        gost = pub.gost_string or pub.title or ''
        ws.append([idx, pub_type_label, pub.title or '', gost])

    ws.column_dimensions['A'].width = 5
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 35
    ws.column_dimensions['D'].width = 60

    bytes_stream = BytesIO()
    wb.save(bytes_stream)
    bytes_stream.seek(0)

    return send_file(
        bytes_stream,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f'publications_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    )


@rating_bp.route('/conferences/export/excel')
@jwt_required()
def export_conferences_excel():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    query = Conference.query.filter(filter_by_role(Conference, user)).filter(Conference.status == 'active')

    author_ids = request.args.getlist('author_ids', type=int)
    if author_ids:
        query = query.filter(Conference.user_id.in_(author_ids))

    search_query = request.args.get('search_query', '')
    if search_query:
        query = query.filter(or_(
            Conference.name.ilike(f'%{search_query}%'),
            Conference.paper_title.ilike(f'%{search_query}%')
        ))

    date_range = request.args.get('date_range', 'all')
    date_from, date_to = parse_date_args(request.args.get('date_from'), request.args.get('date_to'))
    query = apply_date_filter(query, Conference.conference_date, date_range, date_from, date_to)

    conferences = query.all()
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Конференции"
    
    headers = ["ID", "Название", "Роль", "Доклад", "Дата", "Место", "Соавторы"]
    ws.append(headers)
    
    header_fill = PatternFill(start_color="FFC107", end_color="FFC107", fill_type="solid")
    header_font = Font(bold=True, color="000000")
    
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
    
    for conf in conferences:
        ws.append([
            conf.id,
            conf.name,
            conf.role,
            conf.paper_title or "",
            conf.conference_date.strftime("%d.%m.%Y"),
            conf.location or "",
            conf.coauthors or ""
        ])
    
    ws.column_dimensions['A'].width = 5
    ws.column_dimensions['B'].width = 30
    ws.column_dimensions['C'].width = 15
    ws.column_dimensions['D'].width = 30
    ws.column_dimensions['E'].width = 12
    ws.column_dimensions['F'].width = 20
    ws.column_dimensions['G'].width = 30
    
    bytes_stream = BytesIO()
    wb.save(bytes_stream)
    bytes_stream.seek(0)
    
    return send_file(
        bytes_stream,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f"conferences_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    )


# ==================== ЭКСПОРТ В WORD ====================

@rating_bp.route('/awards/export/word')
@jwt_required()
def export_awards_word():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    query = Award.query.filter(filter_by_role(Award, user)).filter(Award.status == 'active')

    author_ids = request.args.getlist('author_ids', type=int)
    if author_ids:
        query = query.filter(Award.user_id.in_(author_ids))

    search_query = request.args.get('search_query', '')
    if search_query:
        query = query.filter(or_(
            Award.title.ilike(f'%{search_query}%'),
            Award.description.ilike(f'%{search_query}%')
        ))

    date_range = request.args.get('date_range', 'all')
    date_from, date_to = parse_date_args(request.args.get('date_from'), request.args.get('date_to'))
    query = apply_date_filter(query, Award.date_received, date_range, date_from, date_to)

    awards = query.all()
    
    doc = Document()
    
    title = doc.add_heading('Отчет о наградах', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_heading('Общая информация', level=2)
    info_table = doc.add_table(rows=1, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_table.rows[0].cells[0].text = 'Всего наград:'
    info_table.rows[0].cells[1].text = str(len(awards))

    doc.add_paragraph()
    
    doc.add_heading('Список наград', level=2)
    
    awards_table = doc.add_table(rows=1, cols=5)
    awards_table.style = 'Light Grid Accent 1'
    
    header_cells = awards_table.rows[0].cells
    header_cells[0].text = 'Название'
    header_cells[1].text = 'Тип'
    header_cells[2].text = 'Организация'
    header_cells[3].text = 'Дата'
    header_cells[4].text = 'Уровень'
    
    for award in awards:
        row_cells = awards_table.add_row().cells
        row_cells[0].text = award.title
        row_cells[1].text = award.award_type
        row_cells[2].text = award.issuer or '-'
        row_cells[3].text = award.date_received.strftime("%d.%m.%Y")
        row_cells[4].text = award.level or ''
    
    bytes_stream = BytesIO()
    doc.save(bytes_stream)
    bytes_stream.seek(0)
    
    return send_file(
        bytes_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"awards_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )


@rating_bp.route('/publications/export/word', methods=['GET'])
@jwt_required()
def export_publications_word():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    pub_type = request.args.get('pub_type', '')
    index_type = request.args.get('index_type', '')
    search_query = request.args.get('search_query', '')
    date_range = request.args.get('date_range', 'all')
    date_from, date_to = parse_date_args(request.args.get('date_from'), request.args.get('date_to'))
    author_ids = request.args.getlist('author_ids', type=int)
    pub_type_map_rev = {v: k for k, v in PUB_MODELS.items()}

    models_to_query = list(PUB_MODELS.values()) if not pub_type else [PUB_MODELS[pub_type]]
    all_publications = []

    for model in models_to_query:
        entity_type = pub_type_map_rev[model]
        query = model.query.filter(filter_by_role(model, user, coauthor_entity_type=entity_type))
        if author_ids:
            query = query.filter(model.user_id.in_(author_ids))

        if index_type == 'scopus' and hasattr(model, 'scopus'):
            query = query.filter(model.scopus == True)
        elif index_type == 'vak' and hasattr(model, 'vak'):
            query = query.filter(model.vak == True)
        elif index_type == 'none' and hasattr(model, 'scopus') and hasattr(model, 'vak'):
            query = query.filter(model.scopus == False, model.vak == False)

        if search_query:
            author_fields = []
            if hasattr(model, 'authors'):
                author_fields.append(model.authors.ilike(f'%{search_query}%'))
            if hasattr(model, 'author_single'):
                author_fields.append(model.author_single.ilike(f'%{search_query}%'))
            if author_fields:
                query = query.filter(or_(model.title.ilike(f'%{search_query}%'), or_(*author_fields)))
            else:
                query = query.filter(model.title.ilike(f'%{search_query}%'))

        if date_range != 'all' and hasattr(model, 'publication_date'):
            if date_range == '3':
                cutoff = datetime.now() - timedelta(days=90)
                query = query.filter(model.publication_date >= cutoff.date())
            elif date_range == '6':
                cutoff = datetime.now() - timedelta(days=180)
                query = query.filter(model.publication_date >= cutoff.date())
            elif date_range == '12':
                cutoff = datetime.now() - timedelta(days=365)
                query = query.filter(model.publication_date >= cutoff.date())
            elif date_range == '36':
                cutoff = datetime.now() - timedelta(days=1095)
                query = query.filter(model.publication_date >= cutoff.date())
            elif date_range == '60':
                cutoff = datetime.now() - timedelta(days=1825)
                query = query.filter(model.publication_date >= cutoff.date())
            elif date_range == 'academic_year':
                now = datetime.now()
                if now.month >= 9:
                    start = date(now.year, 9, 1)
                    end = date(now.year + 1, 8, 31)
                else:
                    start = date(now.year - 1, 9, 1)
                    end = date(now.year, 8, 31)
                query = query.filter(model.publication_date >= start, model.publication_date <= end)
            elif date_range == 'custom':
                if date_from:
                    query = query.filter(model.publication_date >= date_from)
                if date_to:
                    query = query.filter(model.publication_date <= date_to)

        all_publications.extend(query.all())

    if date_range != 'all':
        cutoff_date = None
        if date_range == '3':
            cutoff_date = datetime.now() - timedelta(days=90)
        elif date_range == '6':
            cutoff_date = datetime.now() - timedelta(days=180)
        elif date_range == '12':
            cutoff_date = datetime.now() - timedelta(days=365)
        elif date_range == '36':
            cutoff_date = datetime.now() - timedelta(days=1095)
        elif date_range == '60':
            cutoff_date = datetime.now() - timedelta(days=1825)
        elif date_range == 'academic_year':
            now = datetime.now()
            if now.month >= 9:
                cutoff_date = date(now.year, 9, 1)
            else:
                cutoff_date = date(now.year - 1, 9, 1)

        filtered = []
        for pub in all_publications:
            if pub.publication_date is None:
                continue
            if cutoff_date and pub.publication_date < cutoff_date.date():
                continue
            if date_range == 'academic_year':
                now = datetime.now()
                if now.month >= 9:
                    end = date(now.year + 1, 8, 31)
                else:
                    end = date(now.year, 8, 31)
                if pub.publication_date > end:
                    continue
            if date_range == 'custom' and date_from and pub.publication_date < date_from:
                continue
            if date_range == 'custom' and date_to and pub.publication_date > date_to:
                continue
            filtered.append(pub)
        all_publications = filtered

    doc = Document()

    title = doc.add_heading('Публикации', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    if all_publications:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        header_cells[0].text = '№'
        header_cells[1].text = 'Тип'
        header_cells[2].text = 'Название'
        header_cells[3].text = 'ГОСТ-строка'

        for idx, pub in enumerate(all_publications, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)

            pub_type_label = ''
            for key, model in PUB_MODELS.items():
                if isinstance(pub, model):
                    pub_type_label = PUB_TYPE_LABELS.get(key, key)
                    break
            row_cells[1].text = pub_type_label

            if hasattr(pub, 'url') and pub.url:
                para = row_cells[2].paragraphs[0]
                run = para.add_run(pub.title or 'Публикация')
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.font.underline = True
                para.add_run(f' ({pub.url})')
            else:
                row_cells[2].text = pub.title or ''

            row_cells[3].text = pub.gost_string or pub.title or ''
    else:
        doc.add_paragraph('Публикаций не найдено.')

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'publications_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    )


@rating_bp.route('/conferences/export/word')
@jwt_required()
def export_conferences_word():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    query = Conference.query.filter(filter_by_role(Conference, user)).filter(Conference.status == 'active')

    author_ids = request.args.getlist('author_ids', type=int)
    if author_ids:
        query = query.filter(Conference.user_id.in_(author_ids))

    search_query = request.args.get('search_query', '')
    if search_query:
        query = query.filter(or_(
            Conference.name.ilike(f'%{search_query}%'),
            Conference.paper_title.ilike(f'%{search_query}%')
        ))

    date_range = request.args.get('date_range', 'all')
    date_from, date_to = parse_date_args(request.args.get('date_from'), request.args.get('date_to'))
    query = apply_date_filter(query, Conference.conference_date, date_range, date_from, date_to)

    conferences = query.all()
    
    doc = Document()
    
    title = doc.add_heading('Отчет о конференциях', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_heading('Общая информация', level=2)
    info_table = doc.add_table(rows=1, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_table.rows[0].cells[0].text = 'Всего конференций:'
    info_table.rows[0].cells[1].text = str(len(conferences))
    
    doc.add_paragraph()
    
    doc.add_heading('Список конференций', level=2)
    
    conf_table = doc.add_table(rows=1, cols=5)
    conf_table.style = 'Light Grid Accent 1'
    
    header_cells = conf_table.rows[0].cells
    header_cells[0].text = 'Название'
    header_cells[1].text = 'Роль'
    header_cells[2].text = 'Доклад'
    header_cells[3].text = 'Дата'
    header_cells[4].text = 'Соавторы'
    
    for conf in conferences:
        row_cells = conf_table.add_row().cells
        row_cells[0].text = conf.name
        row_cells[1].text = conf.role
        row_cells[2].text = conf.paper_title or '-'
        row_cells[3].text = conf.conference_date.strftime("%d.%m.%Y")
        row_cells[4].text = conf.coauthors or '-'
    
    bytes_stream = BytesIO()
    doc.save(bytes_stream)
    bytes_stream.seek(0)
    
    return send_file(
        bytes_stream,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f"conferences_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    )


# ==================== ЭКСПОРТ ПОВЫШЕНИЙ КВАЛИФИКАЦИЙ ====================

@rating_bp.route('/qualifications/export/excel', methods=['GET'])
@jwt_required()
def export_qualifications_excel():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    query = Training.query.filter(filter_by_role(Training, user)).filter(Training.status == 'active')

    state_filter = request.args.get('state_issued', '')
    if state_filter == 'yes':
        query = query.filter(Training.state_issued == True)
    elif state_filter == 'no':
        query = query.filter(Training.state_issued == False)

    author_ids = request.args.getlist('author_ids', type=int)
    if author_ids:
        query = query.filter(Training.user_id.in_(author_ids))

    search_query = request.args.get('search_query', '')

    if search_query:
        query = query.filter(or_(
            Training.title.ilike(f'%{search_query}%'),
            Training.organization.ilike(f'%{search_query}%')
        ))

    date_range = request.args.get('date_range', 'all')
    date_from, date_to = parse_date_args(request.args.get('date_from'), request.args.get('date_to'))
    query = apply_date_filter(query, Training.end_date, date_range, date_from, date_to)

    trainings = query.all()
    
    wb = Workbook()
    ws = wb.active
    ws.title = "Тренинги"
    
    headers = ['Дата начала', 'Дата окончания', 'Название курса', 'Организатор', 'Город', "Уровень", 'Часы', 'Номер сертификата', 'Гос. образца']
    ws.append(headers)
    
    for training in trainings:
        ws.append([
            training.start_date.strftime('%d.%m.%Y') if training.start_date else '',
            training.end_date.strftime('%d.%m.%Y') if training.end_date else '',
            training.title or '',
            training.organization or '',
            training.city or '',
            training.level or '',
            training.duration_hours or 0,
            training.certificate_number or '',
            'Да' if training.state_issued else 'Нет'
        ])
    
    for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
        for cell in row:
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    
    ws.column_dimensions['A'].width = 15
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 30
    ws.column_dimensions['D'].width = 25
    ws.column_dimensions['E'].width = 15
    ws.column_dimensions['F'].width = 15
    ws.column_dimensions['G'].width = 12
    ws.column_dimensions['H'].width = 20
    ws.column_dimensions['I'].width = 14
    
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=f'trainings_{user_id}_{datetime.now().strftime("%Y%m%d")}.xlsx'
    )


@rating_bp.route('/qualifications/export/word', methods=['GET'])
@jwt_required()
def export_qualifications_word():
    user_id = get_jwt_identity()
    user = User.query.get_or_404(user_id)
    query = Training.query.filter(filter_by_role(Training, user)).filter(Training.status == 'active')

    state_filter = request.args.get('state_issued', '')
    if state_filter == 'yes':
        query = query.filter(Training.state_issued == True)
    elif state_filter == 'no':
        query = query.filter(Training.state_issued == False)

    author_ids = request.args.getlist('author_ids', type=int)
    if author_ids:
        query = query.filter(Training.user_id.in_(author_ids))

    search_query = request.args.get('search_query', '')
    if search_query:
        query = query.filter(or_(
            Training.title.ilike(f'%{search_query}%'),
            Training.organization.ilike(f'%{search_query}%')
        ))

    date_range = request.args.get('date_range', 'all')
    date_from, date_to = parse_date_args(request.args.get('date_from'), request.args.get('date_to'))
    query = apply_date_filter(query, Training.end_date, date_range, date_from, date_to)

    trainings = query.all()
    
    doc = Document()
    doc.add_heading('Тренинги и повышение квалификации', 0)
    
    if not trainings:
        doc.add_paragraph('Нет данных для экспорта.')
    else:
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        headers = ['Дата начала', 'Дата окончания', 'Название курса', 'Организатор', 'Город', "Уровень", 'Часы', 'Сертификат', 'Гос. образца']
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        for training in trainings:
            row_cells = table.add_row().cells
            row_cells[0].text = training.start_date.strftime('%d.%m.%Y') if training.start_date else ''
            row_cells[1].text = training.end_date.strftime('%d.%m.%Y') if training.end_date else ''
            row_cells[2].text = training.title or ''
            row_cells[3].text = training.organization or ''
            row_cells[4].text = training.city or ''
            row_cells[5].text = training.level or ''
            row_cells[6].text = str(training.duration_hours) if training.duration_hours else '0'
            row_cells[7].text = training.certificate_number or ''
            row_cells[8].text = 'Да' if training.state_issued else 'Нет'
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=f'trainings_{user_id}_{datetime.now().strftime("%Y%m%d")}.docx'
    )
